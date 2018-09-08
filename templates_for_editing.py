#!/usr/bin/env python3
# -*- coding: utf-8 -*-
__description__ = \
    """
So editing html is hard. This should make it better...
NB. Written for python 3, not tested under 2.
"""
__author__ = "Matteo Ferla. [Github](https://github.com/matteoferla)"
__email__ = "matteo.ferla@gmail.com"
__date__ = ""
__license__ = "Cite me!"
__version__ = "1.0"

import os
from bs4 import BeautifulSoup
from docx import Document


def parse_inwards():
    document = Document()
    for file in os.listdir('pyramidstarter/templates'):
        if file.find('.pt') != -1:
            soup=BeautifulSoup(open(os.path.join('pyramidstarter/templates',file)).read(),'html.parser')
            document.add_heading('File: {}'.format(file.replace('.pt', '')), 0)  # title
            for element in soup.find_all(recursive=True):
                if element.has_attr('title'):
                    document.add_heading(element.name, 1)
                    document.add_heading('tooltip', 2)
                    document.add_paragraph(element['title'])
                    document.add_heading('text', 2)
                    document.add_paragraph(element.string)
    document.save('tooltips.docx')




if __name__ == "__main__":
    parse_inwards()